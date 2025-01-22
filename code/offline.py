import os
import sys
import torch
from transformers import pipeline
import pdfplumber
import docx
from concurrent.futures import ThreadPoolExecutor

# Define model and device (CUDA, CPU, or MPS for Apple Silicon)
model_name = "facebook/bart-large-cnn"
device = 0 if torch.cuda.is_available() else ("mps" if torch.backends.mps.is_built() else -1)  # Use GPU if available, MPS for Apple M1/M2, else CPU

# Load the summarizer once, based on platform and device
summarizer = pipeline("summarization", model=model_name, device=device)

# def load_model():
#     """Load the summarizer model in a separate thread to avoid blocking the UI."""
#     global summarizer
#     summarizer = pipeline("summarization", model=model_name, device=device)

# def summarize_chunk(chunk):
#     """Summarizes a single chunk of text."""
#     try:
#         return summarizer(chunk)[0]['summary_text']
#     except Exception as e:
#         print(f"Error summarizing chunk: {e}")
#         return ""

# def summarize_text_incrementally_generator(text):
#     """Generates summarized chunks incrementally."""
#     words = text.split()
#     max_chunk_size = 500  # Adjust chunk size for performance

#     # Split text into chunks of max_chunk_size words
#     chunks = [" ".join(words[i:i+max_chunk_size]) for i in range(0, len(words), max_chunk_size)]
#     for chunk in chunks:
#         yield summarize_chunk(chunk)

# def summarize_file_incrementally(file_path):
#     """Summarizes a file incrementally."""
#     file_type = file_path.split(".")[-1].lower()
#     text = ""
#     if file_type == "pdf":
#         with pdfplumber.open(file_path) as pdf:
#             for page in pdf.pages:
#                 text += page.extract_text()
#     elif file_type == "docx":
#         doc = docx.Document(file_path)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text
#     elif file_type == "txt":
#         with open(file_path, "r", encoding="utf-8") as file:
#             text = file.read()
#     else:
#         return "Unsupported file format."
#     # Generate summarized chunks incrementally
#     return summarize_text_incrementally_generator(text)

def summarize_chunk(chunk):
    """Summarizes a single chunk of text."""
    try:
        max_input_length = 1024  # Maximum input length for the model
        chunk_words = chunk.split()
        if len(chunk_words) > max_input_length:
            sub_chunks = [" ".join(chunk_words[i:i+len(chunk_words)//2]) for i in range(0, len(chunk_words), len(chunk_words)//2)]
            summaries = [summarizer(sub_chunk)[0]['summary_text'] for sub_chunk in sub_chunks]
            return " ".join(summaries)
        else:
            return summarizer(chunk)[0]['summary_text']
    except Exception as e:
        print(f"Error summarizing chunk: {e}")
        return ""

def summarize_text_incrementally_generator(text):
    """Generates summarized chunks incrementally."""
    words = text.split()
    max_chunk_size = 400  # Adjust chunk size for performance
    # print("words", len(words))
    # Split text into chunks of max_chunk_size words
    chunks = [" ".join(words[i:i+max_chunk_size]) for i in range(0, len(words), max_chunk_size)]

    # Process chunks using ThreadPoolExecutor
    with ThreadPoolExecutor() as executor:
        results = executor.map(summarize_chunk, chunks)
        for result in results:
            yield result  # Yield each summarized chunk as it's completed

def summarize_file_incrementally(file_path):
    """Generates summaries incrementally from a file (PDF, DOCX, TXT)."""
    file_type = file_path.split('.')[-1].lower()

    try:
        # Handle PDF files
        if file_type == 'pdf':
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        yield from summarize_text_incrementally_generator(page_text)

        # Handle DOCX files
        elif file_type == 'docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                para_text = para.text
                if para_text:
                    yield from summarize_text_incrementally_generator(para_text)

        # Handle TXT files
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                while True:
                    chunk = f.read(2048)  # Read in chunks of 2048 characters
                    if not chunk:
                        break
                    yield from summarize_text_incrementally_generator(chunk)

        else:
            yield "Unsupported file format."

    except Exception as e:
        print(f"Error reading file: {e}")
        yield f"Error reading file: {e}"

def summarize_folder(folder_path):
    summaries = {}
    for filename in os.listdir(folder_path):
        print(f"Summarizing {filename}...")
        file_path = os.path.join(folder_path, filename)
        for summary in summarize_file_incrementally(file_path):
            print(summary)
            summaries[filename] = summaries.get(filename, []) + [summary]
    return summaries

if __name__ == "__main__":
    article='''This report presents the development of an offline text summarization application tailored for defense-related contexts. Built using Python, PyQt, and transformer-based models (such as Pegasus), the solution condenses large volumes of text—PDFs, Word documents, and plain text—into concise summaries while retaining essential information. Through a streamlined graphical user interface, defense personnel can effortlessly load multiple files or an entire folder, initiate the summarization process, and view incremental results in real-time, thus improving efficiency in data analysis and decision-making.
Text summarization is a critical task in natural language processing, aimed at condensing lengthy textual content into concise, meaningful summaries while preserving the core information. This project presents a robust text summarization application leveraging state-of-the-art AI models. The application supports summarizing various file formats, including PDFs, Word documents, and plain text files, using NLP models like T5-small, Falcon, and Face book’s language models.
The application, built in Python with a user-friendly PyQt5 interface, operates offline and requires no external dependencies. Users can summarize single files or entire folders and customize the summary length (small, medium, or large) based on their needs. Advanced grammar correction ensures that the summaries are polished and professionally formatted.
Key features include an adaptive approach to handling diverse file types, incremental chunk-based summarization for improved performance, and an option to output generated summaries as PDF documents. The system addresses challenges such as managing large documents, optimizing resource usage, and ensuring secure data handling through offline processing. Practical testing demonstrated the software’s effectiveness in significantly reducing manual reading time while still preserving pertinent details vital to defense operations. This offline text summarization solution represents a valuable contribution to the broader effort of applying AI-driven techniques to improve information-processing workflows in sensitive and high-stakes environments.
To enhance usability, the project includes a splash screen with a loading animation, ensuring a smooth user experience while the application initializes. Progress bars provide real-time feedback during the summarization process, and the results can be saved as PDF files with formatted headers. The entire pipeline is designed for efficiency, offering incremental updates to users as the summarization progresses.
The final product demonstrates the integration of cutting-edge AI models with practical software development, showcasing the potential of text summarization technology in various domains such as education, research, and professional documentation.

'''
    summaries = summarize_text_incrementally_generator(article)
    for summary in summaries:
        print(summary)
    # folder_path = "."  # Update this path to a valid directory
    # summaries = summarize_folder(folder_path)
    # for filename, summary_list in summaries.items():
    #     print(f"\nSummaries for {filename}:")
    #     for summary in summary_list:
    #         print(summary)
