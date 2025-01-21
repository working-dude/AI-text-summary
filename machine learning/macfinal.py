import torch
from transformers import pipeline
import pdfplumber
import docx
from concurrent.futures import ThreadPoolExecutor
import os
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QFileDialog, QLabel, QVBoxLayout, QWidget, QTextEdit

model_name = "t5-small"
device = 0 if torch.cuda.is_available() else ("mps" if torch.has_mps else -1)  # Use GPU if available, MPS for Apple M1/M2, else CPU

# Load the summarizer once, based on platform and device
summarizer = pipeline("summarization", model=model_name, device=device)

def summarize_chunk(chunk):
    try:
        return summarizer(chunk)[0]['summary_text']
    except Exception as e:
        print(f"Error summarizing chunk: {e}")
        return ""
    
def summarize_text_incrementally(text):
    words = text.split()
    max_chunk_size = 400  # Larger chunks to speed up

    chunks = [" ".join(words[i:i+max_chunk_size]) for i in range(0, len(words), max_chunk_size)]
    summaries = []
    with ThreadPoolExecutor() as executor:
        summaries = list(executor.map(summarize_chunk, chunks))
    return summaries

def summarize_text_incrementally_generator(text):
    words = text.split()
    max_chunk_size = 400  # Larger chunks to speed up

    chunks = [" ".join(words[i:i+max_chunk_size]) for i in range(0, len(words), max_chunk_size)]

    # Use ThreadPoolExecutor to process chunks in parallel
    with ThreadPoolExecutor() as executor:
        results = executor.map(summarize_chunk, chunks)
        for result in results:
            yield result  # Yield each chunk as it is summarized

def summarize_file_all_text(file_path):
    file_type = file_path.split('.')[-1].lower()
    text = ""

    try:
        # Handle PDF files
        if file_type == 'pdf':
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text

        # Handle DOCX files
        elif file_type == 'docx':
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

        # Handle TXT files
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

        else:
            return "Unsupported file format."

    except Exception as e:
        print(f"Error reading file: {e}")
        return f"Error reading file: {e}"

    if not text.strip():
        return "The document is empty or could not be read."
    
    return summarize_text_incrementally(text)

def summarize_file_incrementally(file_path):
    file_type = file_path.split('.')[-1].lower()
    text = ""

    try:
        # Handle PDF files
        if file_type == 'pdf':
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        yield from summarize_text_incrementally_generator(page_text)  # Process each page incrementally

        # Handle DOCX files
        elif file_type == 'docx':
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            yield from summarize_text_incrementally_generator(text)

        # Handle TXT files
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            yield from summarize_text_incrementally_generator(text)

        else:
            yield "Unsupported file format."

    except Exception as e:
        print(f"Error reading file: {e}")
        yield f"Error reading file: {e}"

def on_click1():
    options = QFileDialog.Options()
    file_path, _ = QFileDialog.getOpenFileName(window, "Open File", "", "All Files (*);;PDF Files (*.pdf);;Word Files (*.docx);;Text Files (*.txt)", options=options)
    # file_path = "ml.pdf"
    if file_path:
        file_type = os.path.splitext(file_path)[1]
        label.setText(f"File Type: {file_type}")
        summary = summarize_file_incrementally(file_path)
        summary_text_edit.setPlainText('\n'.join(summary))
        print("the end")  

def on_click():
    options = QFileDialog.Options()
    file_path, _ = QFileDialog.getOpenFileName(window, "Open File", "", "All Files (*);;PDF Files (*.pdf);;Word Files (*.docx);;Text Files (*.txt)", options=options)
    # file_path = "ml.pdf"
    if file_path:
        file_type = os.path.splitext(file_path)[1]
        label.setText(f"File Type: {file_type}")
        summary = summarize_file_all_text(file_path)
        summary_text_edit.setPlainText('\n'.join(summary))  
        print("the end")

app = QApplication([])
window = QMainWindow()
window.setWindowTitle("PyQt5 GUI")

central_widget = QWidget()
layout = QVBoxLayout(central_widget)

button = QPushButton("Upload File")
button.clicked.connect(on_click)
layout.addWidget(button)

label = QLabel("File Type: ")
layout.addWidget(label)

summary_label = QLabel("Summary: ")
layout.addWidget(summary_label)

summary_text_edit = QTextEdit()
summary_text_edit.setReadOnly(True)
# summary_text_edit.setPlainText(summary)  
layout.addWidget(summary_text_edit)

window.setCentralWidget(central_widget)
window.setGeometry(100, 100, 700, 700)

window.show()
app.exec_()
