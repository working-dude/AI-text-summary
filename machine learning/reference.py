import torch
from transformers import pipeline
import pdfplumber
import docx
from concurrent.futures import ThreadPoolExecutor
import os
import time

device = torch.device("mps")  # Metal Performance Shaders
summarizer = pipeline("summarization", model="./sshleifer/distilbart-xsum-12-6")

def summarize_chunk(chunk):
    try:
        return summarizer(chunk)[0]['summary_text']
    except Exception as e:
        print(f"Error summarizing chunk: {e}")
        return ""

def summarize_text_incrementally(text):
    words = text.split()
    max_chunk_size = 400  # Larger chunks to speed up

    chunks = [" ".join(words[i:i+max_chunk_size]) for i in range(0, len(words), max_chunk_size)]

    # Use ThreadPoolExecutor to process chunks in parallel
    with ThreadPoolExecutor() as executor:
        results = executor.map(summarize_chunk, chunks)
        for result in results:
            yield result  # Yield each chunk as it is summarized

def summarize_file_incrementally(file_path):
    file_type = file_path.split('.')[-1].lower()
    text = ""

    try:
        # Handle PDF files
        if file_type == 'pdf':
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text

        # Handle DOCX files
        elif file_type == 'docx':
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

        # Handle TXT files
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

        else:
            return "Unsupported file format."

    except Exception as e:
        print(f"Error reading file: {e}")
        return f"Error reading file: {e}"

    if not text.strip():
        return "The document is empty or could not be read."
    
    return summarize_text_incrementally(text)

# Start measuring time
start_time = time.time()

# Output the summary incrementally
with open('output.txt', 'w') as file:
    for summary_chunk in summarize_file_incrementally("ronaldo.pdf"):
        # Print each summarized chunk incrementally
        print(summary_chunk)
        
        # Write each chunk to file incrementally
        file.write(summary_chunk + "\n")

# End measuring time
end_time = time.time()
execution_time = end_time - start_time

# Print the execution time
print(f"Execution time: {execution_time} seconds")

# Open the output.txt file after changes are done
os.system('open output.txt')



# ---------- using one page at a time for summary ----------



# import torch
# from transformers import pipeline
# import pdfplumber
# import docx
# from concurrent.futures import ThreadPoolExecutor
# import os
# import time

# device = torch.device("mps")  # Metal Performance Shaders
# summarizer = pipeline("summarization", model="./sshleifer/distilbart-xsum-12-6")

# def summarize_chunk(chunk):
#     try:
#         return summarizer(chunk)[0]['summary_text']
#     except Exception as e:
#         print(f"Error summarizing chunk: {e}")
#         return ""

# def summarize_text_incrementally(text):
#     words = text.split()
#     max_chunk_size = 400  # Larger chunks to speed up

#     chunks = [" ".join(words[i:i+max_chunk_size]) for i in range(0, len(words), max_chunk_size)]

#     # Use ThreadPoolExecutor to process chunks in parallel
#     with ThreadPoolExecutor() as executor:
#         results = executor.map(summarize_chunk, chunks)
#         for result in results:
#             yield result  # Yield each chunk as it is summarized

# def summarize_file_incrementally(file_path):
#     file_type = file_path.split('.')[-1].lower()
#     text = ""

#     try:
#         # Handle PDF files
#         if file_type == 'pdf':
#             with pdfplumber.open(file_path) as pdf:
#                 for page in pdf.pages:
#                     page_text = page.extract_text()
#                     if page_text:
#                         yield from summarize_text_incrementally(page_text)  # Process each page incrementally

#         # Handle DOCX files
#         elif file_type == 'docx':
#             doc = docx.Document(file_path)
#             text = "\n".join([para.text for para in doc.paragraphs])
#             yield from summarize_text_incrementally(text)

#         # Handle TXT files
#         elif file_type == 'txt':
#             with open(file_path, 'r', encoding='utf-8') as f:
#                 text = f.read()
#             yield from summarize_text_incrementally(text)

#         else:
#             yield "Unsupported file format."

#     except Exception as e:
#         print(f"Error reading file: {e}")
#         yield f"Error reading file: {e}"

# # Start measuring time
# start_time = time.time()

# # Output the summary incrementally
# with open('output.txt', 'w') as file:
#     for summary_chunk in summarize_file_incrementally("ronaldo.pdf"):
#         if summary_chunk != "":
#         # Print each summarized chunk incrementally
#             print(summary_chunk)
            
#             # Write each chunk to file incrementally
#             file.write(summary_chunk + "\n")

# # End measuring time
# end_time = time.time()
# execution_time = end_time - start_time

# # Print the execution time
# print(f"Execution time: {execution_time} seconds")

# # Open the output.txt file after changes are done
# os.system('open output.txt')
